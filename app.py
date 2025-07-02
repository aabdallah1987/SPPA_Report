import streamlit as st
import sqlite3
import datetime
import os
from docx import Document
from io import BytesIO

# --- Constants & Configuration ---
DB_FILE = "lpi_sessions.db"

TASKS = {
    1: {"base": [
            "Simple WH-Question Exchange", "Simple WH-Question Exchange", 
            "Simple WH-Question Exchange", 
            "Role-play (no complication)", "Asking Basic Questions"
        ], "probe_level": 2},
    2: {"base": [
            "Narration (Past)", "Detailed Description", "Instructions", 
            "Reporting on current events", "Role-play (minor complication)", "Narration (Future)",
            "Narration (Present)"
        ], "probe_level": 3}
}

TASK_PROBES = {
    2: ["Narration (Past)", "Detailed Description", "Instructions", "Role-play (minor complication)", "Reporting on current events"],
    3: ["Abstract Discussion", "Supporting Opinion", "Hypothesizing"]
}

TASK_INSTRUCTIONS = {
    "Simple WH-Question Exchange": "Ask the learner 5-6 simple WH-questions about a topic from the warm-up (when, where, with whom, how long, etc.).",
    "Asking Basic Questions": "Have the learner ask you 5 questions about a common object or topic (e.g., your phone, your car, your weekend).",
    "Role-play (no complication)": "Read a short, straightforward role-play scenario in English, then play it out in the target language. There should be no problems or complications (e.g., successfully ordering a coffee).",
    "Narration (Past)": "Ask the learner to tell a story about a personal experience—something they did or witnessed—from beginning to end in chronological order with detail. This should ideally be in the past time frame.",
    "Narration (Present)": "Ask the learner about something they do routinely in detail, from start to finish, in a step-by-step, chronological order (e.g., a daily work routine, exercise plan, or a typical weekend day).",
    "Narration (Future)": "Ask about the learner's future plans for something specific (e.g., an upcoming trip, project, or celebration).",
    "Detailed Description": "Ask the learner to describe something from their life in detail (e.g., their home, a room, a building, a vehicle).",
    "Instructions": "Ask the learner to explain how to do something—like applying for a job, buying a ticket, or cooking a meal.",
    "Reporting on current events": "Ask the learner to report 5-6 facts about a recent event or something widely discussed (e.g., gas prices, a recent game, a news story).",
    "Role-play (minor complication)": "Read a short role-play scenario in English, then play it out in the target language. Introduce a small problem (e.g., a hotel has no rooms available).",
    "Abstract Discussion": "Introduce a societal issue (e.g., media influence, technology use). Ask the learner to discuss the topic abstractly — what people say, general perspectives.",
    "Supporting Opinion": "Present a controversial issue with two sides. Ask the learner which side they support and why. They should express and justify their opinion clearly.",
    "Hypothesizing": "Introduce a scenario and ask the learner to speculate about the consequences. For example, ‘What would happen if...?’ Use open-ended hypothetical prompts."
}

SURVEY_QUESTIONS = [
    "Where are you now? (City, state.)",
    "What do you do? Military? How long have you been working in the military?",
    "What did you do before the military? Did you do any jobs before or during college and high school?",
    "Where do you live? Do you live in a house or a condo? Do you own or rent? How long have you lived there? Who lives with you?",
    "Do you have pets or any animals?",
    "Where are you from originally?",
    "Did you go to college? Where and what did you study?",
    "What do you like to do in your free time? (leisure activities or any other activities like hobbies, workout, spending time with friends and family, etc.)",
    "Where did you travel in or out of the US in the past few years for other than work? How long did you spend there?"
]

RATINGS = ["🔴 Total Breakdown", "🟠 Partial Response", "🟡 Minimal Sustained Response", "🟢 Fully Sustained Response"]
RATING_MAP = {rating: i for i, rating in enumerate(RATINGS)}


# --- Database Logic ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            session_id INTEGER PRIMARY KEY AUTOINCREMENT,
            learner_name TEXT NOT NULL, learner_id TEXT NOT NULL, language TEXT NOT NULL,
            interview_date TIMESTAMP NOT NULL, initial_level INTEGER NOT NULL, final_level TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS task_responses (
            response_id INTEGER PRIMARY KEY AUTOINCREMENT, session_id INTEGER NOT NULL,
            task_name TEXT NOT NULL, task_level INTEGER NOT NULL, rating TEXT NOT NULL,
            FOREIGN KEY (session_id) REFERENCES sessions (session_id)
        )
    """)
    conn.commit()
    conn.close()

def save_session_to_db(session_data):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    initial_level_to_save = session_data.get('initial_level', session_data.get('current_level'))
    # MODIFIED: Use the selected interview_date and pass blank strings for name/id
    cursor.execute("""
        INSERT INTO sessions (learner_name, learner_id, language, interview_date, initial_level, final_level)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (
        "", "", session_data['language'],
        session_data['interview_date'], initial_level_to_save, session_data['final_level']
    ))
    session_id = cursor.lastrowid
    for task in session_data['tasks_completed']:
        cursor.execute("INSERT INTO task_responses (session_id, task_name, task_level, rating) VALUES (?, ?, ?, ?)",
                       (session_id, task['name'], task['level'], task['rating']))
    conn.commit()
    conn.close()

# --- Scoring Logic ---
def calculate_final_level(current_level, tasks):
    base_tasks = [t for t in tasks if t['level'] == current_level]
    probe_tasks = [t for t in tasks if t['level'] > current_level]

    base_green_count = sum(1 for t in base_tasks if RATING_MAP.get(t['rating']) == 3)
    base_yellow_or_better_count = sum(1 for t in base_tasks if RATING_MAP.get(t['rating'], -1) >= 2)
    base_orange_or_better_count = sum(1 for t in base_tasks if RATING_MAP.get(t['rating'], -1) >= 1)
    
    probe_yellow_or_better_count = sum(1 for t in probe_tasks if RATING_MAP.get(t['rating'], -1) >= 2)
    probe_orange_count = sum(1 for t in probe_tasks if RATING_MAP.get(t['rating']) == 1)
    
    all_probes_are_breakdown = False
    if probe_tasks:
        all_probes_are_breakdown = all(RATING_MAP.get(t['rating']) == 0 for t in probe_tasks)

    is_strong_base = base_green_count >= 5
    is_minimal_base = base_yellow_or_better_count >= 4

    if current_level == 1:
        # MODIFIED: Specific check for strong base + all probe breakdown
        if is_strong_base and all_probes_are_breakdown:
            return "1", "Strong base performance but breakdown on all probe tasks."
        if all_probes_are_breakdown:
            if base_orange_or_better_count >= 3:
                return "0+", "Sustained partial performance at base with no performance on probes."
            else:
                return "0", "Breakdown at base and on probes."
        
        if is_strong_base and probe_yellow_or_better_count >= 2: return "MOVE_TO_L2", "Performance indicates learner may be at Level 2."
        if is_strong_base and probe_orange_count >= 1: return "1+", "Strong base with partial probes."
        if is_minimal_base: return "1", "Minimal sustained performance at base."
        
    if current_level == 2:
        if is_strong_base and probe_yellow_or_better_count >= 2: return "MOVE_TO_L3", "Performance indicates learner may be at Level 3."
        if is_strong_base and probe_orange_count >= 1: return "2+", "Strong base with partial probes."
        if is_minimal_base: return "2", "Minimal sustained performance at base."
        
    if current_level == 3:
        l3_tasks = [t for t in tasks if t['level'] == 3]
        l3_yellow_or_better_count = sum(1 for t in l3_tasks if RATING_MAP.get(t['rating'], -1) >= 2)
        if l3_yellow_or_better_count >= 2:
            return "3", "Sustained performance at Level 3."
        else:
            return "2+", "Could not sustain performance at Level 3."

    return "INVALID", "Invalid Test. The combination of ratings did not produce a valid score. Please retest the student with tasks appropriate for the current and upper levels."

# --- Main App Logic ---
def initialize_state(force_reset=False):
    if force_reset:
        for key in list(st.session_state.keys()):
            del st.session_state[key]
    if 'page' not in st.session_state:
        st.session_state.page = 'welcome'
        st.session_state.language = ''
        st.session_state.interview_date = datetime.date.today()
        st.session_state.initial_level = None; st.session_state.current_level = None
        st.session_state.tasks_completed = []; st.session_state.final_level = None;
        st.session_state.final_reasoning = ""

def set_page(page_name):
    st.session_state.page = page_name

def generate_task_list(level):
    if level == 1:
        return [
            {'name': "Simple WH-Question Exchange", 'level': 1, 'status': 'pending'},
            {'name': "Narration (Past)", 'level': 2, 'status': 'pending'},
            {'name': "Simple WH-Question Exchange", 'level': 1, 'status': 'pending'},
            {'name': "Detailed Description", 'level': 2, 'status': 'pending'},
            {'name': "Simple WH-Question Exchange", 'level': 1, 'status': 'pending'},
            {'name': "Instructions", 'level': 2, 'status': 'pending'},
            {'name': "Role-play (no complication)", 'level': 1, 'status': 'pending'},
            {'name': "Reporting on current events", 'level': 2, 'status': 'pending'},
            {'name': "Asking Basic Questions", 'level': 1, 'status': 'pending'}
        ]

    if level == 3:
        l3_tasks = TASK_PROBES[level][:]
        return [{'name': task, 'level': level, 'status': 'pending'} for task in l3_tasks]

    base_tasks_list = TASKS[level]['base'][:]
    probe_level = TASKS[level]['probe_level']
    probe_tasks_list = TASK_PROBES[probe_level][:]
    
    last_task_name = "Asking Basic Questions" if level == 1 else "Future Narration"
    last_task_obj = None
    if last_task_name in base_tasks_list:
        base_tasks_list.remove(last_task_name)
        last_task_obj = {'name': last_task_name, 'level': level, 'status': 'pending'}

    interleaved_tasks = []
    while base_tasks_list or probe_tasks_list:
        if base_tasks_list: interleaved_tasks.append({'name': base_tasks_list.pop(0), 'level': level, 'status': 'pending'})
        if probe_tasks_list: interleaved_tasks.append({'name': probe_tasks_list.pop(0), 'level': probe_level, 'status': 'pending'})

    if last_task_obj: interleaved_tasks.append(last_task_obj)
    return interleaved_tasks

# --- Helper function to generate .docx report ---
def create_docx_report():
    document = Document()
    document.add_heading('Speaking Proficiency Placement Assessment (SPPA) Report', 0)

    document.add_heading('Session Details', level=1)
    # MODIFIED: Use interview_date instead of name/id
    document.add_paragraph(f"Interview Date: {st.session_state.interview_date.strftime('%B %d, %Y')}")
    document.add_paragraph(f"Language: {st.session_state.language}")
    
    document.add_heading('Assessment Result', level=1)
    p = document.add_paragraph(); p.add_run('Suggested Proficiency Level: ').bold = True
    p.add_run(st.session_state.final_level)
    p = document.add_paragraph(); p.add_run('Reasoning: ').bold = True
    p.add_run(st.session_state.final_reasoning)

    document.add_heading('Task Performance Breakdown', level=1)
    for i, task in enumerate(st.session_state.tasks_completed):
        document.add_paragraph(
            f"Task {i+1}: {task['name']} (Level {task['level']}) — Rating: {task['rating']}", 
            style='List Bullet'
        )
        topic = st.session_state.get(f"topic_{i}", ""); prompt = st.session_state.get(f"prompt_{i}", "")
        comment = st.session_state.get(f"comment_{i}", "")
        document.add_paragraph(f"\tTopic: {topic}"); document.add_paragraph(f"\tTask Prompt: {prompt}")
        document.add_paragraph(f"\tComment on Learner Response: {comment}")

    buffer = BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer

# --- UI Rendering Functions ---
def render_welcome_page():
    st.title("Speaking Proficiency Placement Assessment (SPPA) Tool")
    st.write("---")
    st.header("Student Info: Enter Session Information")
    
    # MODIFIED: Removed name/id inputs, added date input
    lang = st.text_input("Language Being Assessed", key="lang_input")
    date = st.date_input("Interview Date", key="date_input")
    
    if st.button("Begin Session", disabled=not (lang and date)):
        st.session_state.language = lang
        st.session_state.interview_date = date
        set_page('survey'); st.rerun()

def render_survey_page():
    st.title("Phase 1: Language Survey (Warm-Up) – 10 Minutes")
    st.info("Ask the following background questions to build rapport and estimate the learner's initial working level. Feel free to ask more questions as needed.")
    st.warning("Note on Sensitive Topics: Avoid asking detailed questions about military work, and refrain from inquiring about family matters unless the learner initiates the conversation about it.")
    for q in SURVEY_QUESTIONS: st.markdown(f"- {q}")
    st.write("---")
    st.header("Step 1: Select Initial Working Level")
    
    col1, col2 = st.columns(2)
    if col1.button("Level 1 (Intermediate-Low/High)"):
        st.session_state.initial_level = 1; st.session_state.current_level = 1
        st.session_state.tasks_to_do = generate_task_list(1)
        set_page('tasks'); st.rerun()
    if col2.button("Level 2 (Advanced-Low/High)"):
        st.session_state.initial_level = 2; st.session_state.current_level = 2
        st.session_state.tasks_to_do = generate_task_list(2)
        set_page('tasks'); st.rerun()
            
def render_tasks_page():
    level = st.session_state.current_level
    if not level:
        st.error("No level selected."); st.button("Back to Welcome", on_click=set_page, args=['welcome']); return

    if level == 2: st.markdown(f"<h1 style='color: green;'>Phase 2: Task Phase (Testing at Level {level})</h1>", unsafe_allow_html=True)
    elif level == 3: st.markdown(f"<h1 style='color: blue;'>Phase 2: Task Phase (Testing at Level {level})</h1>", unsafe_allow_html=True)
    else: st.title(f"Phase 2: Task Phase (Testing at Level {level})")
    
    for i, task in enumerate(st.session_state.tasks_to_do):
        task_key = f"task_{i}_level_{level}"; st.subheader(f"Task {i+1}: {task['name']} (Level {task['level']} Function)")
        if task['name'] in TASK_INSTRUCTIONS: st.caption(f"💡 Instruction: {TASK_INSTRUCTIONS[task['name']]}")

        cols = st.columns(len(RATINGS))
        for j, rating in enumerate(RATINGS):
            if cols[j].button(rating, key=f"{task_key}_rating_{j}", use_container_width=True):
                task['rating'] = rating; task['status'] = 'done'; st.rerun()
        
        if task.get('status') == 'done': st.success(f"Selected: {task['rating']}")
        st.write("---")
    
    num_tasks_done = sum(1 for t in st.session_state.tasks_to_do if t.get('status') == 'done')
    all_tasks_rated = num_tasks_done == len(st.session_state.tasks_to_do)

    if st.button("Calculate Score", type="primary", disabled=not all_tasks_rated):
        if all(t.get('rating') == RATINGS[0] for t in st.session_state.tasks_to_do):
            st.error("Retest the learner with different level-appropriate questions or rate the learner at Level 0.")
            if st.button("Clear All and Restart"): initialize_state(force_reset=True); st.rerun()
        else:
            st.session_state.tasks_completed.extend(st.session_state.tasks_to_do)
            st.session_state.tasks_to_do = []
            level, reason = calculate_final_level(st.session_state.current_level, st.session_state.tasks_completed)
            
            if level == "INVALID": st.error(reason)
            elif "MOVE_TO" in level:
                new_level = int(level[-1]); st.session_state.current_level = new_level
                st.session_state.tasks_to_do = generate_task_list(new_level)
                st.info(f"Performance warrants moving to Level {new_level}. Presenting new tasks."); st.rerun()
            else:
                st.session_state.final_level = level; st.session_state.final_reasoning = reason
                save_session_to_db(st.session_state)
                set_page('summary'); st.rerun()

def render_summary_page():
    st.title("Final Proficiency Assessment Summary")
    if not st.session_state.get('balloons_shown', False):
        st.balloons(); st.session_state.balloons_shown = True
    
    st.header("Session Details")
    # MODIFIED: Display interview_date instead of name/id
    st.write(f"**Interview Date:** {st.session_state.interview_date.strftime('%B %d, %Y')}")
    st.write(f"**Language:** {st.session_state.language}")
    
    st.header("Assessment Result")
    st.metric(label="Suggested Proficiency Level", value=st.session_state.final_level)
    st.caption(f"Reasoning: {st.session_state.final_reasoning}")
    st.header("Task Performance Breakdown")
    
    for i, task in enumerate(st.session_state.tasks_completed):
        st.markdown(f"**Task {i+1}: {task['name']} (Level {task['level']}) — Rating: {task['rating']}**")
        topic_key = f"topic_{i}"; prompt_key = f"prompt_{i}"; comment_key = f"comment_{i}"
        st.text_input("Topic:", key=topic_key); st.text_area("Task Prompt:", key=prompt_key)
        st.text_area("Comment on Learner Response:", key=comment_key)
        st.write("---")

    st.write("") # Spacer
    report_buffer = create_docx_report()
    st.download_button(
        label="Download Report as .docx",
        data=report_buffer,
        file_name=f"SPPA_Report_{st.session_state.interview_date}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if st.button("Start New Interview"):
        initialize_state(force_reset=True); st.rerun()

# --- Main Application Router ---
if not os.path.exists(DB_FILE): init_db()
initialize_state()

if st.session_state.page == 'welcome': render_welcome_page()
elif st.session_state.page == 'survey': render_survey_page()
elif st.session_state.page == 'tasks': render_tasks_page()
elif st.session_state.page == 'summary': render_summary_page()
else: set_page('welcome'); st.rerun()
